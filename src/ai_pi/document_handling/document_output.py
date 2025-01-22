"""
This file takes the output from reviewer_agent.py and creates a Word doc
given the original document and the review output, yielding a new doc which
looks as if a reviewer left comments and suggested revisions. The idea is to
leave the user able to parse through each and address them individually using
this output.
"""
import docx
from docx.oxml import OxmlElement
from fuzzywuzzy import fuzz
import json

def enable_track_changes(doc):
    """Enable track changes in the document."""
    # Create track revisions tag
    tag = OxmlElement('w:trackRevisions')
    
    # Get the settings element
    settings = doc.settings._element
    settings.append(tag)

def add_high_level_review(doc, high_level_review):
    """Add high-level review as a new section at the start of the document."""
    # Add title
    doc.add_heading('Scientific Review Summary', level=1)
    
    # Add metrics section if available
    if 'metrics' in high_level_review:
        doc.add_heading('Review Metrics', level=2)
        metrics = high_level_review['metrics']
        if isinstance(metrics, str):
            doc.add_paragraph(metrics)
        elif isinstance(metrics, dict):
            # Handle structured metrics
            for category, scores in metrics.items():
                doc.add_paragraph(f"{category.title()}:", style='Heading 3')
                if isinstance(scores, dict):
                    for subcategory, score in scores.items():
                        doc.add_paragraph(f"- {subcategory.title()}: {score}", style='List Bullet')
                else:
                    doc.add_paragraph(f"Score: {scores}")
    
    # Add overall assessment
    doc.add_heading('Overall Assessment', level=2)
    doc.add_paragraph(high_level_review.get('overall_assessment', 'No overall assessment provided.'))
    
    # Add key strengths
    doc.add_heading('Key Strengths', level=2)
    strengths = high_level_review.get('key_strengths', [])
    if strengths:
        for strength in strengths:
            doc.add_paragraph(strength.strip(), style='List Bullet')
    else:
        doc.add_paragraph('No key strengths specified.')
    
    # Add key weaknesses
    doc.add_heading('Key Weaknesses', level=2)
    weaknesses = high_level_review.get('key_weaknesses', [])
    if weaknesses:
        for weakness in weaknesses:
            doc.add_paragraph(weakness.strip(), style='List Bullet')
    else:
        doc.add_paragraph('No key weaknesses specified.')
    
    # Add recommendations
    doc.add_heading('Key Recommendations', level=2)
    recommendations = high_level_review.get('recommendations', [])
    if recommendations:
        for rec in recommendations:
            doc.add_paragraph(rec.strip(), style='List Bullet')
    else:
        doc.add_paragraph('No recommendations provided.')
    
    # Add communication review section
    if 'communication_review' in high_level_review:
        doc.add_heading('Writing and Communication Review', level=2)
        comm_review = high_level_review['communication_review']
        
        # Writing assessment (as regular paragraph)
        doc.add_paragraph(comm_review.get('writing_assessment', 'No writing assessment provided.'))
        
        # Narrative strengths
        if comm_review.get('narrative_strengths'):
            doc.add_heading('Narrative Strengths', level=3)
            for strength in comm_review['narrative_strengths']:
                clean_text = strength.strip().lstrip('•').strip()
                doc.add_paragraph(clean_text, style='List Bullet')
        
        # Areas for improvement
        if comm_review.get('narrative_weaknesses'):
            doc.add_heading('Areas for Improvement', level=3)
            for weakness in comm_review['narrative_weaknesses']:
                clean_text = weakness.strip().lstrip('•').strip()
                doc.add_paragraph(clean_text, style='List Bullet')
        
        # Style recommendations
        if comm_review.get('style_recommendations'):
            doc.add_heading('Style Recommendations', level=3)
            for rec in comm_review['style_recommendations']:
                clean_text = rec.strip().lstrip('•').strip()
                doc.add_paragraph(clean_text, style='List Bullet')
    
    # Add page break before the original document
    doc.add_page_break()

def output_commented_document(input_doc_path, review_struct, output_doc_path, match_threshold=90, verbose=False):
    """Process the document and add AI review comments and suggestions."""
    # Create a copy of the original document
    doc = docx.Document(input_doc_path)
    doc.save(output_doc_path)
    doc = docx.Document(output_doc_path)
    
    # Update how we extract the high-level review
    if 'reviews' in review_struct:
        high_level_review = {
            'metrics': review_struct['reviews'].get('metrics', 'No metrics provided.'),
            'overall_assessment': review_struct['reviews'].get('main_review', {}).get('overall_assessment', ''),
            'key_strengths': review_struct['reviews'].get('main_review', {}).get('key_strengths', '').split('\n') 
                if isinstance(review_struct['reviews'].get('main_review', {}).get('key_strengths'), str) else [],
            'key_weaknesses': review_struct['reviews'].get('main_review', {}).get('key_weaknesses', '').split('\n')
                if isinstance(review_struct['reviews'].get('main_review', {}).get('key_weaknesses'), str) else [],
            'recommendations': review_struct['reviews'].get('main_review', {}).get('global_suggestions', '').split('\n')
                if isinstance(review_struct['reviews'].get('main_review', {}).get('global_suggestions'), str) else []
        }
        
        # Create a new document for the review section
        review_doc = docx.Document()
        add_high_level_review(review_doc, high_level_review)
        
        # Insert review section at the beginning of original document
        for element in reversed(review_doc.element.body):
            doc.element.body.insert(0, element)

    enable_track_changes(doc)
    
    print(f"Processing document with {len(doc.paragraphs)} paragraphs")
    
    # Flatten section reviews into lists
    all_match_strings = []
    all_comments = []
    all_revisions = []
    
    # Process review items from section reviews
    print(f"\nProcessing review items:")
    
    # Process top-level review items
    if 'review_items' in review_struct:
        for item in review_struct['review_items']:
            if isinstance(item, dict):
                all_match_strings.append(item.get('match_string', ''))
                all_comments.append(item.get('comment', ''))
                all_revisions.append(item.get('revision', ''))

    # Add revisions from the revisions section
    for revision in review_struct.get('revisions', []):
        if 'original_text' in revision:
            all_match_strings.append(revision['original_text'])
            all_comments.append(revision.get('comment', ''))
            all_revisions.append(revision.get('new_text', ''))

    print("\nLooking for these matches:", all_match_strings)
    
    matches_found = 0
    # Keep original matches in a list that won't be modified
    all_matches = list(zip(all_match_strings, all_comments, all_revisions))
    
    # Track which matches have been successfully processed
    processed_matches = set()
    
    # Process each paragraph looking for matches
    for paragraph in doc.paragraphs:
        text = paragraph.text
        normalized_text = ' '.join(text.split())
        
        # Try each match that hasn't been processed yet
        for match, comment, revision in all_matches:
            if match in processed_matches:
                continue
                
            normalized_match = ' '.join(match.split())
            
            # Try exact match first
            if normalized_match in normalized_text:
                match_ratio = 100
                match_location = normalized_text.index(normalized_match)
            else:
                # Use fuzzy matching as fallback
                match_ratio = fuzz.token_set_ratio(normalized_match, normalized_text)
                if match_ratio < match_threshold:
                    continue
                    
                # Find best position for fuzzy match
                match_location = normalized_text.find(normalized_match)
                if match_location == -1:
                    # Use approximate position if exact substring not found
                    words = normalized_text.split()
                    match_words = normalized_match.split()
                    for i in range(len(words)):
                        if fuzz.ratio(words[i], match_words[0]) > match_threshold:
                            match_location = normalized_text.find(words[i])
                            break
            
            try:
                # Split and process the paragraph
                before_match = normalized_text[:match_location]
                after_match = normalized_text[match_location + len(normalized_match):]
                
                # Clear and rebuild paragraph
                paragraph.clear()
                
                if before_match:
                    paragraph.add_run(before_match)
                
                # Add matched text with comment/revision
                if revision and revision.strip():
                    # Add deletion with comment
                    del_run = paragraph.add_run(normalized_match)
                    del_run.font.strike = True
                    del_run.add_comment(f"{comment} (Match confidence: {match_ratio}%)", author="AIPI", initials="AI")
                    
                    # Add revision as new text
                    ins_run = paragraph.add_run(f" {revision} ")
                    ins_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
                else:
                    # Just add comment
                    match_run = paragraph.add_run(normalized_match)
                    match_run.add_comment(comment, author="AIPI", initials="AI")
                
                if after_match:
                    paragraph.add_run(after_match)
                
                processed_matches.add(match)
                matches_found += 1
                
            except Exception as e:
                print(f"Error processing match '{match}': {str(e)}")
                continue
    
    # Report unmatched strings
    unmatched = set(m[0] for m in all_matches) - processed_matches
    if unmatched:
        print("\nWarning: The following matches were not found in the document:")
        for match in unmatched:
            print(f"- '{match}'")
    
    print(f"\nTotal matches found: {matches_found}")
    print(f"Saving document to {output_doc_path}")
    doc.save(output_doc_path)

if __name__ == "__main__":
    # Test data with high-level review
    # Load sample review data from JSON
    import json
    
    with open("/home/christian/projects/agents/ai_pi/processed_documents/ScolioticFEPaper_v7_20250113_030430/ScolioticFEPaper_v7_reviewed.json", "r") as f:
        review_data = json.load(f)
    
    # Test paths - adjust these to your actual file locations
    input_path = "examples/ScolioticFEPaper_v7.docx"
    output_path = "examples/test_output_with_review.docx"
    
    try:
        # Pass the entire review_data structure directly
        output_commented_document(input_path, review_data, output_path, verbose=True)
        print(f"Successfully created reviewed document with high-level summary at {output_path}")
    except Exception as e:
        print(f"Error processing document: {str(e)}")
