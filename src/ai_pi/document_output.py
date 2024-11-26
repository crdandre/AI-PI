"""
This file takes the output from reviewer_agent.py and creates a Word doc
given the original document and the review output, yielding a new doc which
looks as if a reviewer left comments and suggested revisions. The idea is to
leave the user able to parse through each and address them individually using
this output.
"""
import docx
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def enable_track_changes(doc):
    """Enable track changes in the document."""
    # Create track revisions tag
    tag = OxmlElement('w:trackRevisions')
    
    # Get the settings element
    settings = doc.settings._element
    settings.append(tag)

def output_commented_document(input_doc_path, document_review_items, output_doc_path):
    """Process the document and add AI review comments and suggestions."""
    doc = docx.Document(input_doc_path)
    
    # Enable track changes
    enable_track_changes(doc)
    
    print(f"Processing document with {len(doc.paragraphs)} paragraphs")
    print("\nLooking for these matches:", document_review_items["match_strings"])
    
    matches_found = 0
    # Iterate through each paragraph in the document
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        print(f"\nProcessing paragraph {i}: {text[:50]}...")
        
        for match, comment, revision in zip(
            document_review_items["match_strings"],
            document_review_items["comments"],
            document_review_items["revisions"]
        ):
            if match in text:
                matches_found += 1
                print(f"Found match: '{match}'")
                try:
                    # Find the start and end position of the matched text
                    start_idx = text.index(match)
                    end_idx = start_idx + len(match)
                    
                    # Split the paragraph into parts
                    before_match = text[:start_idx]
                    after_match = text[end_idx:]
                    
                    # Clear the paragraph's content
                    paragraph.clear()
                    
                    # Add text before the match
                    if before_match:
                        paragraph.add_run(before_match)
                    
                    # Add the matched text with comment and revision
                    if revision:
                        # Create deletion run and add comment to it
                        del_run = paragraph.add_run()
                        
                        # Create run properties
                        rPr = OxmlElement('w:rPr')
                        
                        # Create deletion element
                        del_element = OxmlElement('w:del')
                        del_element.set(qn('w:author'), 'AIPI')
                        del_element.set(qn('w:date'), '2024-03-21T12:00:00Z')
                        
                        # Create text element for deletion
                        del_text = OxmlElement('w:t')
                        del_text.set(qn('xml:space'), 'preserve')
                        del_text.text = match
                        
                        # Build the XML structure
                        del_run._element.append(rPr)
                        del_run._element.append(del_element)
                        del_element.append(del_text)
                        
                        # Add comment to the deletion run
                        del_run.add_comment(comment, author="AIPI", initials="AI")
                        
                        # Create insertion run (without comment)
                        ins_run = paragraph.add_run()
                        
                        # Create run properties
                        rPr = OxmlElement('w:rPr')
                        
                        # Create insertion element
                        ins_element = OxmlElement('w:ins')
                        ins_element.set(qn('w:author'), 'AIPI')
                        ins_element.set(qn('w:date'), '2024-03-21T12:00:00Z')
                        
                        # Create text element for insertion
                        ins_text = OxmlElement('w:t')
                        ins_text.set(qn('xml:space'), 'preserve')
                        ins_text.text = revision
                        
                        # Build the XML structure
                        ins_run._element.append(rPr)
                        ins_run._element.append(ins_element)
                        ins_element.append(ins_text)
                    else:
                        # If no revision, just add the matched text with comment
                        match_run = paragraph.add_run(match)
                        match_run.add_comment(comment, author="AIPI", initials="AI")
                    
                    # Add text after the match
                    if after_match:
                        paragraph.add_run(after_match)
                    
                    print(f"Added comment: '{comment}'")
                    if revision:
                        print(f"Added revision: '{revision}'")
                        
                except Exception as e:
                    print(f"Error processing revision: {str(e)}")
                    print(f"Error type: {type(e)}")
                    continue
    
    print(f"\nTotal matches found: {matches_found}")
    print(f"Saving document to {output_doc_path}")
    doc.save(output_doc_path)

if __name__ == "__main__":
    # Test data
    test_review_items = {
        "match_strings": ["Maximizing the accuracy of material property", "The pelvis was constrained"],
        "comments": ["This is a sample comment", "Another review comment"],
        "revisions": ["suggested revision", "another suggestion"]
    }
    
    # Test paths - adjust these to your actual file locations
    input_path = "/home/christian/projects/agents/ai_pi/examples/ScolioticFEPaper_v7.docx"
    output_path = "/home/christian/projects/agents/ai_pi/examples/test_output.docx"
    
    try:
        output_commented_document(input_path, test_review_items, output_path)
        print(f"Successfully created reviewed document at {output_path}")
    except Exception as e:
        print(f"Error processing document: {str(e)}")

