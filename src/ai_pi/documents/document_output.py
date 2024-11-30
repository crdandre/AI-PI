"""
This file takes the output from reviewer_agent.py and creates a Word doc
given the original document and the review output, yielding a new doc which
looks as if a reviewer left comments and suggested revisions. The idea is to
leave the user able to parse through each and address them individually using
this output.
"""
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fuzzywuzzy import fuzz
from fuzzywuzzy import process

def enable_track_changes(doc):
    """Enable track changes in the document."""
    # Create track revisions tag
    tag = OxmlElement('w:trackRevisions')
    
    # Get the settings element
    settings = doc.settings._element
    settings.append(tag)

def output_commented_document(input_doc_path, document_review_items, output_doc_path, match_threshold=90, verbose=False):
    """Process the document and add AI review comments and suggestions."""
    doc = docx.Document(input_doc_path)
    enable_track_changes(doc)
    
    print(f"Processing document with {len(doc.paragraphs)} paragraphs")
    print("\nLooking for these matches:", document_review_items["match_strings"])
    
    matches_found = 0
    # Keep original matches in a list that won't be modified
    all_matches = list(zip(
        document_review_items["match_strings"],
        document_review_items["comments"],
        document_review_items["revisions"]
    ))
    
    # Track which matches have been successfully processed
    processed_matches = set()
    
    # Iterate through each paragraph in the document
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        normalized_text = ' '.join(text.split())
        
        # Try each match that hasn't been processed yet
        for match, comment, revision in all_matches:
            # Skip if we've already processed this match
            if match in processed_matches:
                continue
                
            normalized_match = ' '.join(match.split())
            
            # Try exact match first with more context
            full_context = normalized_text
            if normalized_match in full_context:
                match_ratio = 100
                match_location = full_context.index(normalized_match)
                text = full_context
                match = normalized_match
                if verbose:
                    print(f"Exact match found for: '{match}'")
            else:
                # Enhanced fuzzy matching with better context handling
                match_ratio = fuzz.token_set_ratio(normalized_match, full_context)
                if match_ratio >= match_threshold:
                    try:
                        # Use larger context windows for better matching
                        window_size = len(normalized_match) * 2
                        windows = []
                        
                        for i in range(max(0, len(text) - window_size + 1)):
                            window_text = text[i:i + window_size]
                            window_score = fuzz.token_set_ratio(normalized_match, window_text)
                            windows.append((window_text, i, window_score))
                        
                        # Sort windows by score and get the best match
                        best_window = max(windows, key=lambda x: x[2])
                        match_location = best_window[1]
                        
                        # Verify the match position
                        context_before = text[max(0, match_location-50):match_location]
                        context_after = text[match_location:match_location+len(normalized_match)+50]
                        
                        if verbose:
                            print(f"Context before: '{context_before}'")
                            print(f"Matched text: '{text[match_location:match_location+len(normalized_match)]}'")
                            print(f"Context after: '{context_after}'")
                            
                    except Exception as e:
                        if verbose:
                            print(f"Error in fuzzy matching: {str(e)}")
                        continue
                else:
                    continue
            
            try:
                # Split the paragraph into parts
                before_match = text[:match_location]
                after_match = text[match_location + len(match):]
                
                # Clear the paragraph's content
                paragraph.clear()
                
                # Add text before the match
                if before_match:
                    paragraph.add_run(before_match)
                
                # Add the matched text with comment and revision
                if revision:
                    # Create deletion run and add comment to it
                    del_run = paragraph.add_run()
                    
                    # Create run properties
                    rPr = OxmlElement('w:rPr')
                    
                    # Create deletion element
                    del_element = OxmlElement('w:del')
                    del_element.set(qn('w:author'), 'AIPI')
                    del_element.set(qn('w:date'), '2024-03-21T12:00:00Z')
                    
                    # Create text element for deletion
                    del_text = OxmlElement('w:t')
                    del_text.set(qn('xml:space'), 'preserve')
                    del_text.text = match
                    
                    # Build the XML structure
                    del_run._element.append(rPr)
                    del_run._element.append(del_element)
                    del_element.append(del_text)
                    
                    # Add comment to the deletion run
                    del_run.add_comment(f"{comment} (Match confidence: {match_ratio}%)", author="AIPI", initials="AI")
                    
                    # Create insertion run
                    ins_run = paragraph.add_run()
                    rPr = OxmlElement('w:rPr')
                    ins_element = OxmlElement('w:ins')
                    ins_element.set(qn('w:author'), 'AIPI')
                    ins_element.set(qn('w:date'), '2024-03-21T12:00:00Z')
                    ins_text = OxmlElement('w:t')
                    ins_text.set(qn('xml:space'), 'preserve')
                    ins_text.text = revision
                    ins_run._element.append(rPr)
                    ins_run._element.append(ins_element)
                    ins_element.append(ins_text)
                else:
                    # If no revision, just add the matched text with comment
                    match_run = paragraph.add_run(match)
                    match_run.add_comment(f"{comment} (Match confidence: {match_ratio}%)", author="AIPI", initials="AI")
                
                # Add text after the match
                if after_match:
                    paragraph.add_run(after_match)
                
                # Mark this match as processed
                processed_matches.add(match)
                matches_found += 1
                
                print(f"Found match: '{match}' with confidence {match_ratio}%")
                print(f"Added comment: '{comment}'")
                if revision:
                    print(f"Added revision: '{revision}'")
                    
            except Exception as e:
                print(f"Error processing match '{match}': {str(e)}")
                continue
    
    # Report unmatched strings
    unmatched = set(m[0] for m in all_matches) - processed_matches
    if unmatched:
        print("\nWarning: The following matches were not found in the document:")
        for match in unmatched:
            print(f"- '{match}'")
    
    print(f"\nTotal matches found: {matches_found}")
    print(f"Saving document to {output_doc_path}")
    doc.save(output_doc_path)

if __name__ == "__main__":
    # Test data
    # test_review_items = {
    #     "match_strings": ["Maximizing the accuracy of material property", "The pelvis was constrained"],
    #     "comments": ["This is a sample comment", "Another review comment"],
    #     "revisions": ["suggested revision", "another suggestion"]
    # }
    
    test_review_items = {
        'match_strings': [
            'The current study will build upon our previously published work on pediatric patient-specific FE modeling and growth', 
            'Vertebral growth was modeled through adaptation of a region-specific orthotropic thermal expansion method described by Balasubramanian', 
            'and adjusted linearly according to Risser', 
        ],
        'revisions': [
            'The current study builds upon our previously published work on pediatric patient-specific FE modeling and growth, which has shown promising results in predicting scoliotic curve progression.',
            'Vertebral growth was modeled using a region-specific orthotropic thermal expansion method, as described by Balasubramanian et al. (reference), which has been validated in previous studies.',
            'The value for was initially set to 0.4 MPa-1, as reported by Shi et al. (reference), and then adjusted linearly based on the Risser sign, with a clear explanation of how the adjustment was made.',
        ], 
        'comments': [
            'Added context to clarify the significance of the current study.', 
            'Added reference and context to support the methodology.', 
            'Improved clarity and added reference to support the methodology.'
        ]
    }
    
    # Test paths - adjust these to your actual file locations
    input_path = "/home/christian/projects/agents/ai_pi/examples/ScolioticFEPaper_v7.docx"
    output_path = "/home/christian/projects/agents/ai_pi/examples/test_output.docx"
    
    try:
        output_commented_document(input_path, test_review_items, output_path)
        print(f"Successfully created reviewed document at {output_path}")
    except Exception as e:
        print(f"Error processing document: {str(e)}")

