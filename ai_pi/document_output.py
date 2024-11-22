"""
This file takes the output from reviewer_agent.py and creates a Word doc
given the original document and the review output, yielding a new doc which
looks as if a reviewer left comments and suggested revisions. The idea is to
leave the user able to parse through each and address them individually using
this output.
"""
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_comment(paragraph, comment_text, author="AI Reviewer"):
    """Add a comment to a paragraph in the Word document."""
    # Create the comment reference
    comment_id = 0  # You might want to track this across multiple comments
    comment_reference = OxmlElement('w:commentReference')
    comment_reference.set(qn('w:id'), str(comment_id))
    
    # Create the actual comment
    comment = OxmlElement('w:comment')
    comment.set(qn('w:id'), str(comment_id))
    comment.set(qn('w:author'), author)
    comment.set(qn('w:date'), '2024-01-01T00:00:00Z')
    
    # Add the comment text
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = comment_text
    r.append(t)
    p.append(r)
    comment.append(p)
    
    # Make sure the document has a comments part
    part = paragraph._parent._parent.part
    if not hasattr(part, 'comments_part'):
        # Create a new comments part
        from docx.opc.constants import CONTENT_TYPE as CT
        from docx.opc.part import Part
        
        # Create the comments part
        comments_part = Part(
            part.package,
            '/word/comments.xml',
            CT.WML_COMMENTS,
            OxmlElement('w:comments')
        )
        part.package.parts.append(comments_part)
        part.comments_part = comments_part
    
    # Add the comment to the document
    part.comments_part.element.append(comment)
    paragraph._element.append(comment_reference)
    
    return comment

def output_commented_document(input_doc_path, document_review_items, output_doc_path):
    """
    Process the document and add AI review comments and suggestions.
    
    Args:
        input_doc_path (str): Path to the input document
        document_review_items (dict): Dictionary containing matches, comments, and revisions
        output_doc_path (str): Path where the reviewed document should be saved
    """
    doc = Document(input_doc_path)
    
    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        text = paragraph.text
        
        # Check each match string
        for match, comment, revision in zip(
            document_review_items["match_strings"],
            document_review_items["comments"],
            document_review_items["revisions"]
        ):
            if match in text:
                # Add the comment
                add_comment(paragraph, comment)
                
                # If there's a revision, show it in red
                if revision:
                    # Create a new run with the suggested revision
                    run = paragraph.add_run(f" [Suggested revision: {revision}]")
                    run.font.color.rgb = RGBColor(255, 0, 0)
    
    # Save the modified document
    doc.save(output_doc_path)

